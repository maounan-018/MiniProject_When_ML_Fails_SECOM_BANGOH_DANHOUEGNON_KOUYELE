"""Generate the complete project report as a .docx file."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE    = Path(r'C:\DAVID DANHOUEGNON ECC\OneDrive\DOC DE COURS 2A\Machine Learning\MiniProject_When_ML_Fails_SECOM')
FIGS    = BASE / 'figures'
REPORTS = BASE / 'reports'
REPORTS.mkdir(exist_ok=True)
OUT     = REPORTS / 'MiniProject_When_ML_Fails_SECOM_BANGOH_DANHOUEGNON_KOUYELE.docx'

# ── Colors ────────────────────────────────────────────────────────────────────
BLUE        = RGBColor(0x2E, 0x75, 0xB6)
DARK_BLUE   = RGBColor(0x1F, 0x49, 0x7D)
ORANGE      = RGBColor(0xC5, 0x5A, 0x11)
GREEN_DARK  = RGBColor(0x37, 0x56, 0x23)
RED         = RGBColor(0xC0, 0x00, 0x00)
BLACK       = RGBColor(0x00, 0x00, 0x00)
TH_BG       = 'D5E8F0'   # table header background (light blue)
BROKEN_BG   = 'FCE4D6'   # broken model row background (light orange)
FIXED_BG    = 'E2EFDA'   # corrected model row background (light green)
REF_BG      = 'EBF3FB'   # reference model row background (light blue-white)
BONUS_BG    = 'FFF2CC'   # bonus section background (light yellow)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=True, bottom=True, left=True, right=True,
                    color='BFBFBF', size='4'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, active in [('top', top), ('bottom', bottom),
                          ('left', left), ('right', right)]:
        if active:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    size)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)

def set_para_border_bottom(para, color='2E75B6', size='12'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_run_font(run, name='Calibri', size_pt=11, bold=False,
                 italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1, color=None, border_bottom=False):
    para = doc.add_paragraph()
    para.style = doc.styles[f'Heading {level}']
    run = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.bold  = True
    run.font.size  = Pt({1:16, 2:13, 3:12}.get(level, 11))
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = DARK_BLUE if level == 1 else BLUE
    if border_bottom:
        set_para_border_bottom(para)
    para.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    para.paragraph_format.space_after  = Pt(6)
    return para

def add_body(doc, text, bold_parts=None, italic=False, indent=False):
    """Add a body paragraph, optionally bolding certain substrings."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(5)
    para.paragraph_format.space_before = Pt(0)
    if indent:
        para.paragraph_format.left_indent = Cm(0.8)
    if bold_parts is None:
        run = para.add_run(text)
        set_run_font(run, italic=italic)
        return para
    # Handle bold_parts: list of (substring, is_bold) — applied in order
    remaining = text
    for chunk, is_bold in bold_parts:
        idx = remaining.find(chunk)
        if idx == -1:
            continue
        if idx > 0:
            r = para.add_run(remaining[:idx])
            set_run_font(r)
        r = para.add_run(chunk)
        set_run_font(r, bold=is_bold, italic=italic)
        remaining = remaining[idx+len(chunk):]
    if remaining:
        r = para.add_run(remaining)
        set_run_font(r)
    return para

def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    set_run_font(run)
    para.paragraph_format.space_after  = Pt(3)
    para.paragraph_format.space_before = Pt(0)
    if level > 0:
        para.paragraph_format.left_indent = Cm(level * 0.8)
    return para

def add_note(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size_pt=9.5, italic=True, color=RGBColor(0x60,0x60,0x60))
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.space_before = Pt(2)
    return para

def add_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run_font(run, size_pt=10, italic=True, color=RGBColor(0x40,0x40,0x40))
    para.paragraph_format.space_after  = Pt(10)
    para.paragraph_format.space_before = Pt(2)
    return para

def add_figure(doc, fname, caption, width_in=5.8):
    path = FIGS / fname
    if not path.exists():
        add_note(doc, f'[Figure not found: {fname}]')
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    run = para.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    add_caption(doc, caption)

def add_table(doc, headers, rows, col_widths_in,
              row_colors=None, bold_col=None):
    """
    headers:      list of str
    rows:         list of list of str
    col_widths_in: list of float (inches)
    row_colors:   list of hex strings (one per data row), or None
    bold_col:     column index to bold
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hrow = table.rows[0]
    for i, (cell, hdr) in enumerate(zip(hrow.cells, headers)):
        cell.width = Inches(col_widths_in[i])
        set_cell_bg(cell, TH_BG)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        set_run_font(run, bold=True, size_pt=10, color=DARK_BLUE)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = (row_colors[ri] if row_colors and ri < len(row_colors) else None)
        for ci, (cell, val) in enumerate(zip(row.cells, row_data)):
            cell.width = Inches(col_widths_in[ci])
            if bg:
                set_cell_bg(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if ci == 0
                           else WD_ALIGN_PARAGRAPH.CENTER)
            is_bold = (bold_col is not None and ci == bold_col)
            run = p.add_run(str(val))
            set_run_font(run, bold=is_bold, size_pt=10)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_separator(doc):
    para = doc.add_paragraph()
    set_para_border_bottom(para, color='D5E8F0', size='6')
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(8)

def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(WD_BREAK.PAGE)

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Header
from docx.oxml import OxmlElement as _oxml
hdr = doc.sections[0].header
hdr_para = hdr.paragraphs[0]
hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hdr_run = hdr_para.add_run('ECC Spring 2026  |  Mini-Project: When Machine Learning Fails')
set_run_font(hdr_run, size_pt=9, italic=True, color=RGBColor(0x80,0x80,0x80))

# Footer with page numbers
from docx.oxml.ns import qn as _qn
ftr = doc.sections[0].footer
ftr_para = ftr.paragraphs[0]
ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
ftr_run = ftr_para.add_run()
set_run_font(ftr_run, size_pt=9, color=RGBColor(0x80,0x80,0x80))
fldChar1 = _oxml('w:fldChar'); fldChar1.set(_qn('w:fldCharType'), 'begin')
instrText = _oxml('w:instrText'); instrText.text = ' PAGE '
fldChar2 = _oxml('w:fldChar'); fldChar2.set(_qn('w:fldCharType'), 'end')
ftr_run._r.append(fldChar1); ftr_run._r.append(instrText); ftr_run._r.append(fldChar2)

# ── COVER ─────────────────────────────────────────────────────────────────────
def cover_line(text, size=11, bold=False, italic=False,
               color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size_pt=size, bold=bold, italic=italic,
                 color=color or RGBColor(0x40, 0x40, 0x40))
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def cover_label(text):
    """Small uppercase section label on the cover."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text.upper())
    set_run_font(r, size_pt=10, bold=True,
                 color=RGBColor(0x70, 0x70, 0x70))
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after  = Pt(6)
    return p

# — Top course info (haut de page)
cover_line('Centrale Casablanca — ECC Spring 2026',
           size=13, italic=True, color=RGBColor(0x55, 0x55, 0x55),
           space_before=10, space_after=4)
cover_line('Introduction to AI and Machine Learning',
           size=13, italic=True, color=RGBColor(0x55, 0x55, 0x55),
           space_before=0, space_after=0)

# — Large vertical gap before title
for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(22)

# — Main title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run('When Machine Learning Fails')
r_title.font.name = 'Calibri'
r_title.font.size = Pt(46)
r_title.font.bold = True
r_title.font.color.rgb = DARK_BLUE
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(16)

# — Subtitle
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run('Diagnosing and Repairing Failure Modes of a Non-Linear Model')
set_run_font(r_sub, size_pt=20, italic=True, color=BLUE)
p_sub.paragraph_format.space_after = Pt(12)
set_para_border_bottom(p_sub, color='2E75B6', size='12')

# — Type + dataset
cover_line('Mini-Project Report', size=15,
           color=RGBColor(0x55, 0x55, 0x55), space_before=10, space_after=5)
cover_line('SECOM Dataset  ·  RandomForestClassifier',
           size=14, italic=True,
           color=RGBColor(0x45, 0x45, 0x45), space_before=0, space_after=0)

# — Authors
cover_label('Authors')
cover_line('David Danhouegnon  ·  Levis Bangoh  ·  Kouyele Samaa Ashley',
           size=20, bold=True, color=DARK_BLUE,
           space_before=0, space_after=0)

# — Professors
cover_label('Professors')
cover_line('Kawtar Zerhouni  &  Rym Nassih',
           size=18, bold=False, color=BLUE,
           space_before=0, space_after=5)
cover_line('UTER MID@S', size=13, italic=True,
           color=RGBColor(0x70, 0x70, 0x70),
           space_before=0, space_after=0)

add_page_break(doc)

# ── TABLE OF CONTENTS ─────────────────────────────────────────────────────────
toc_h = doc.add_paragraph()
toc_h.style = doc.styles['Heading 1']
toc_hr = toc_h.add_run('Table of Contents')
toc_hr.font.name = 'Calibri'; toc_hr.font.bold = True
toc_hr.font.size = Pt(16); toc_hr.font.color.rgb = DARK_BLUE
set_para_border_bottom(toc_h)
toc_h.paragraph_format.space_before = Pt(4)
toc_h.paragraph_format.space_after  = Pt(12)

toc_para = doc.add_paragraph()
toc_run  = toc_para.add_run()
_fc1 = OxmlElement('w:fldChar');  _fc1.set(qn('w:fldCharType'), 'begin')
_it  = OxmlElement('w:instrText'); _it.set(qn('xml:space'), 'preserve')
_it.text = 'TOC \\o "1-2" \\h \\z \\u'
_fc2 = OxmlElement('w:fldChar');  _fc2.set(qn('w:fldCharType'), 'separate')
_fc3 = OxmlElement('w:fldChar');  _fc3.set(qn('w:fldCharType'), 'end')
toc_run._r.append(_fc1); toc_run._r.append(_it)
toc_run._r.append(_fc2); toc_run._r.append(_fc3)

add_note(doc,
    'Right-click on this field → "Update Field" to populate the table of contents.')
add_page_break(doc)

# ── SECTION 1 ─────────────────────────────────────────────────────────────────
add_heading(doc, '1. Research Question and Chosen Dataset', level=1, border_bottom=True)

add_body(doc,
    'The SECOM dataset records 590 process measurements from a semiconductor factory. '
    'Each observation corresponds to one production run, labelled Pass (−1) or Fail (1). '
    'The dataset is appropriate for this project because it is high-dimensional, '
    'severely imbalanced, and contains missing values, which creates conditions where '
    'model failures are easy to hide behind aggregate accuracy.',
)

add_body(doc,
    'Research question:',
    bold_parts=[('Research question:', True)],
)

# Boxed research question
rq = doc.add_paragraph()
rq.paragraph_format.left_indent  = Cm(1.0)
rq.paragraph_format.right_indent = Cm(1.0)
rq.paragraph_format.space_before = Pt(4)
rq.paragraph_format.space_after  = Pt(10)
rq_run = rq.add_run(
    '"Does a Random Forest trained on SECOM rely on a non-causal shortcut feature '
    'strongly correlated with the target in training, such that its ranking performance '
    '(ROC-AUC) collapses when that correlation is reversed at deployment time?"'
)
set_run_font(rq_run, italic=True, size_pt=11, color=DARK_BLUE)

add_body(doc,
    'This question is falsifiable. The hypothesis is rejected if: '
    '(a) shortcut_alarm\'s feature importance remains comparable to that of genuine sensor features, '
    'or (b) ROC-AUC does not degrade significantly when the shortcut is inverted.',
)

add_heading(doc, 'Dataset and Preprocessing', level=2)
add_body(doc,
    'The timestamp is dropped. Features with more than 50% missing values are removed. '
    'The categorical Phase column is one-hot encoded. Constant columns are removed. '
    'Median imputation is wrapped inside the scikit-learn Pipeline to prevent data leakage. '
    'The train/test split (70/30) is stratified on the target.',
)

add_table(doc,
    headers=['Property', 'Value'],
    rows=[
        ['Total observations',              '1,567'],
        ['Raw features',                    '593'],
        ['Fail count / rate',               '104 / 6.64%'],
        ['Features removed (>50% missing)', '28'],
        ['Features removed (constant)',     '116'],
        ['Features after cleaning',         '452'],
        ['Train rows',                      '1,096  (73 Fail — 6.66%)'],
        ['Test rows',                       '471    (31 Fail — 6.58%)'],
    ],
    col_widths_in=[3.2, 3.0],
)
add_note(doc, 'Table 1. Dataset and preprocessing summary.')

# ── SECTION 2 ─────────────────────────────────────────────────────────────────
add_heading(doc, '2. Reference Model and Observed Symptom', level=1, border_bottom=True)

add_body(doc,
    'The reference model is a scikit-learn Pipeline: median imputation followed by '
    'RandomForestClassifier with class_weight="balanced_subsample", '
    'min_samples_leaf=2, 400 trees, and random_state=42. '
    'The stratified train/test split uses the same fixed seed.',
)

add_heading(doc, 'Observed Symptom', level=2)
add_body(doc,
    'On the test set, accuracy reaches 93.4% — but the model predicts Pass for every single '
    'observation. Recall and F1 on the Fail class are both zero. '
    'This is the classic class imbalance trap: the majority-class baseline is already '
    'correct 93.4% of the time, so accuracy is a misleading metric for this problem.',
)
add_body(doc,
    'ROC-AUC = 0.796 on the clean test set confirms that the model has some ranking ability — '
    'it separates Fail from Pass better than random — but the default threshold (0.5) never '
    'crosses the probability of the minority class. '
    'This asymmetry makes ROC-AUC the primary metric for the shortcut experiment.',
)

add_table(doc,
    headers=['Setting', 'Accuracy', 'Balanced Acc.', 'Recall (Fail)', 'F1 (Fail)', 'ROC-AUC'],
    rows=[
        ['Reference — train clean', '1.0000', '1.0000', '1.0000', '1.0000', '1.0000'],
        ['Reference — test clean',  '0.9342', '0.5000', '0.0000', '0.0000', '0.7958'],
    ],
    col_widths_in=[2.4, 0.95, 1.1, 1.15, 0.9, 0.9],
    row_colors=[REF_BG, REF_BG],
)
add_note(doc, 'Table 2. Reference pipeline results. The model predicts no Fail on the test set.')

# ── SECTION 3 ─────────────────────────────────────────────────────────────────
add_heading(doc, '3. Causal Hypothesis and Controlled Experiment', level=1, border_bottom=True)

add_heading(doc, 'Causal Hypothesis', level=2)
add_body(doc,
    'Hypothesis: a sufficiently flexible non-linear model will preferentially rely on any '
    'feature that is reliably predictive of the target in the training data, '
    'regardless of whether that feature is causally related to the outcome. '
    'When its predictive value is removed at test time, the model\'s apparent performance '
    'will collapse in proportion to the degree of reliance on that feature.',
)
add_body(doc,
    'Falsification criterion: if the model does not assign disproportionate importance to the '
    'shortcut feature, or if performance is stable when the shortcut is inverted, '
    'the hypothesis is rejected.',
)

add_heading(doc, 'Controlled Variable: shortcut_alarm', level=2)
add_body(doc,
    'A synthetic binary feature shortcut_alarm is created from the true labels. '
    'In training, it mirrors the target at 98% reliability (2% random flips). '
    'At evaluation time, the same test observations are presented under three conditions — '
    'keeping the shortcut intact changes nothing but the shortcut-target relationship:',
)
add_bullet(doc, 'Aligned (ID):   shortcut_alarm ≈ target  (correlation ≈ +0.88)')
add_bullet(doc, 'Randomised:    shortcut_alarm is random  (correlation ≈  0.00)')
add_bullet(doc, 'Inverted (OOD): shortcut_alarm ≈ ¬target (correlation ≈ −0.87)')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_table(doc,
    headers=['Condition', 'Correlation with target'],
    rows=[
        ['Train — aligned',    '+0.8825'],
        ['Test  — aligned',    '+0.8834'],
        ['Test  — randomised', '+0.0046'],
        ['Test  — inverted',   '−0.8713'],
    ],
    col_widths_in=[2.8, 2.2],
    row_colors=[FIXED_BG, None, None, BROKEN_BG],
)
add_note(doc, 'Table 3. Controlled shortcut diagnostic. Rows, labels, model class, and seed are identical across conditions.')

add_heading(doc, 'Result: ROC-AUC Collapses when Shortcut is Inverted', level=2)
add_body(doc,
    'The broken model — trained with the aligned shortcut — achieves ROC-AUC = 0.993 when '
    'evaluated on the aligned test set. When the same model is evaluated on the inverted test '
    'set (same rows, same labels, only the shortcut reversed), ROC-AUC falls to 0.019. '
    'This is worse than random, which confirms that the model has learned to rely almost '
    'exclusively on shortcut_alarm rather than on the genuine sensor features.',
)

add_figure(doc, 'fig4_confusion_shortcut.png',
    'Figure 1. Confusion matrices on the same test rows: shortcut aligned (left) vs inverted (right). '
    'The model predicts Fail correctly when the shortcut is aligned and predicts incorrectly '
    'when it is inverted — the confusion pattern reverses almost completely.',
    width_in=5.8)

add_heading(doc, 'Sensitivity Analysis: Effect Grows with Shortcut Reliability', level=2)
add_body(doc,
    'To confirm that the effect is proportional to the shortcut\'s apparent predictive value, '
    'we sweep the reliability parameter across four levels (0.55, 0.70, 0.85, 0.98) and '
    'train a separate model at each level. The gap between ID and OOD performance widens '
    'monotonically with reliability, confirming that the model\'s shortcut dependency is driven '
    'by how reliable the feature appears during training.',
)
add_figure(doc, 'fig2_sensitivity.png',
    'Figure 2. ROC-AUC (left) and F1 (right) as a function of shortcut reliability. '
    'The divergence between in-distribution (blue) and out-of-distribution (orange) performance '
    'grows monotonically — the more reliable the shortcut in training, the larger the collapse.',
    width_in=6.0)

add_heading(doc, 'Feature Importance: shortcut_alarm Dominates', level=2)
add_body(doc,
    'Both native (MDI) importance and permutation importance (measured by ROC-AUC drop) '
    'place shortcut_alarm as the single most important feature by a wide margin. '
    'This confirms that the model\'s apparent performance is almost entirely driven by the '
    'non-causal shortcut, not by the 452 genuine sensor features.',
)
add_figure(doc, 'fig3_feature_importance.png',
    'Figure 3. Feature importance for the broken model (shortcut in orange). '
    'Native MDI importance (left) and permutation importance (right) both place shortcut_alarm '
    'far above all 452 sensor features — the model has learned the wrong signal.',
    width_in=6.0)

# ── SECTION 4 ─────────────────────────────────────────────────────────────────
add_heading(doc, '4. Proposed Correction and Evaluation', level=1, border_bottom=True)

add_heading(doc, 'Correction: Remove the Non-Causal Feature', level=2)
add_body(doc,
    'The correction is to remove shortcut_alarm before training the corrected pipeline. '
    'This directly targets the identified cause — feature reliance on a non-causal proxy — '
    'rather than masking the symptom. It is not sufficient to tune the decision threshold, '
    'change class weighting, or improve the evaluation split, because those interventions '
    'do not eliminate the model\'s incentive to rely on the shortcut.',
)
add_body(doc,
    'In the notebook, the broken and corrected pipelines are exposed as independent '
    'functions run_broken_pipeline(...) and run_corrected_pipeline(...), both reproducible '
    'from the same prepared data and fixed seed.',
)

add_heading(doc, 'Before / After Comparison', level=2)
add_body(doc,
    'After removing the shortcut, ROC-AUC returns from 0.019 (OOD, broken) to 0.796, '
    'matching the clean reference model exactly. This confirms the causal diagnosis: '
    'the shortcut was the sole source of the apparent performance advantage. '
    'The gain in OOD robustness is +0.777 ROC-AUC points.',
)
add_body(doc,
    'Note: recall and F1 remain zero across all configurations because the class imbalance '
    '(6.6% Fail) means the 0.5 decision threshold never fires on the minority class. '
    'This is why ROC-AUC — a threshold-free ranking metric — is the appropriate measure '
    'for the shortcut experiment.',
    bold_parts=[('Note:', True)],
)

add_table(doc,
    headers=['Setting', 'Balanced Acc.', 'Recall (Fail)', 'F1 (Fail)', 'ROC-AUC'],
    rows=[
        ['Broken — ID (shortcut aligned)',   '0.5000', '0.0000', '0.0000', '0.9933'],
        ['Broken — OOD (shortcut inverted)', '0.5000', '0.0000', '0.0000', '0.0193'],
        ['Corrected — shortcut removed',     '0.5000', '0.0000', '0.0000', '0.7958'],
        ['Reference — clean test set',       '0.5000', '0.0000', '0.0000', '0.7958'],
    ],
    col_widths_in=[2.7, 1.05, 1.15, 0.9, 0.95],
    row_colors=[BROKEN_BG, BROKEN_BG, FIXED_BG, REF_BG],
    bold_col=4,
)
add_note(doc, 'Table 4. Before/after evaluation. ROC-AUC is the key metric (threshold-free). '
              'Corrected and Reference are identical because shortcut_alarm was the only added feature.')

add_figure(doc, 'fig1_roc_comparison.png',
    'Figure 4. ROC-AUC across evaluation conditions. The broken model collapses from 0.993 '
    'to 0.019 when the shortcut is inverted. The correction restores performance to the '
    'reference level (0.796).',
    width_in=5.5)

# ── SECTION 5 ─────────────────────────────────────────────────────────────────
add_heading(doc, '5. Experimental Validation — Variance Across Seeds', level=1, border_bottom=True)

add_body(doc,
    'A single train/test split could yield a misleading conclusion, particularly with only '
    '31 Fail observations in the test set. We therefore repeat the full experiment on '
    'five random seeds (7, 11, 19, 23, 42) and report mean ± standard deviation.',
)

add_table(doc,
    headers=['Setting', 'ROC-AUC mean', 'ROC-AUC std', 'F1 (Fail) mean', 'F1 (Fail) std'],
    rows=[
        ['Broken — ID (shortcut aligned)',   '0.9928', '0.0011', '0.0000', '0.0000'],
        ['Broken — OOD (shortcut inverted)', '0.0339', '0.0161', '0.0000', '0.0000'],
        ['Corrected — clean test',           '0.7682', '0.0248', '0.0000', '0.0000'],
    ],
    col_widths_in=[2.7, 1.05, 1.05, 1.1, 1.05],
    row_colors=[BROKEN_BG, BROKEN_BG, FIXED_BG],
    bold_col=1,
)
add_note(doc, 'Table 5. Multi-seed validation on 5 independent splits. '
              'The effect is consistent: the gap between broken_id and broken_ood '
              'exceeds 0.95 ROC-AUC points on every seed.')

add_figure(doc, 'fig5_multiseed.png',
    'Figure 5. ROC-AUC mean ± std across 5 seeds. The shortcut collapse (broken OOD) '
    'and the correction gain are both stable — the effect is not due to split variability.',
    width_in=5.2)

add_body(doc,
    'The negligible standard deviation of broken_id (0.0011) confirms that the shortcut '
    'mechanism is deterministic given the high reliability (0.98). '
    'The corrected model shows more variance (0.025) which is expected — the genuine '
    'SECOM signal is weak and noisy, and small changes in the train set composition '
    'affect the ranking ability.',
)

# ── SECTION 6 — BONUS A ───────────────────────────────────────────────────────
add_page_break(doc)

bonus_para = doc.add_paragraph()
bonus_run  = bonus_para.add_run('  BONUS — Failure Mode A  ')
set_run_font(bonus_run, bold=True, size_pt=9.5, color=RGBColor(0x7F,0x60,0x00))
bonus_para.paragraph_format.space_after = Pt(2)

add_heading(doc, '6. Failure Mode A: Class Imbalance Failure', level=1, border_bottom=True)

add_heading(doc, 'Symptom', level=2)
add_body(doc,
    'A Random Forest trained without class weighting achieves 93.4% accuracy on the SECOM '
    'test set while predicting Pass for every observation — recall on the Fail class is '
    'exactly zero. High accuracy masks a model that is completely useless for defect detection.',
)

add_heading(doc, 'Causal Hypothesis and Controlled Test', level=2)
add_body(doc,
    'Hypothesis: because the Fail class represents only 6.6% of samples, a model that '
    'minimises unweighted cross-entropy learns to never predict Fail. '
    'The optimal constant predictor achieves 93.4% accuracy, so the model has no incentive '
    'to learn the minority class at all.',
)
add_body(doc,
    'Test: we train two models on the same data — one without class weighting and one '
    'with class_weight="balanced_subsample". We additionally optimise the decision threshold '
    'on a held-out validation set using the F1 score as criterion.',
)

add_heading(doc, 'Correction: Class Weighting + Threshold Optimisation', level=2)
add_body(doc,
    'The corrected pipeline uses class_weight="balanced_subsample" (each tree resamples the '
    'training set to balance the classes) and optimises the decision threshold on a 25% '
    'validation split. The optimal threshold found on validation is 0.134 — far below the '
    'default 0.5 — confirming that the model\'s calibration is poor under severe imbalance.',
)

add_table(doc,
    headers=['Setting', 'Threshold', 'Balanced Acc.', 'Recall (Fail)', 'F1 (Fail)', 'ROC-AUC'],
    rows=[
        ['Broken — no class weight, test',                      '0.500', '0.5000', '0.0000', '0.0000', '0.7792'],
        ['Corrected — balanced weight, threshold 0.5',          '0.500', '0.5000', '0.0000', '0.0000', '0.7921'],
        ['Corrected — balanced weight, optimised threshold',     '0.134', '0.5695', '0.1935', '0.1967', '0.7921'],
    ],
    col_widths_in=[2.7, 0.8, 1.05, 1.1, 0.9, 0.85],
    row_colors=[BROKEN_BG, FIXED_BG, FIXED_BG],
    bold_col=3,
)
add_note(doc, 'Table 6. Class imbalance results. ROC-AUC is unchanged by threshold (it is threshold-free). '
              'Recall improves from 0% to 19.4% with the optimal threshold.')

add_figure(doc, 'fig6_confusion_imbalance.png',
    'Figure 6. Class imbalance — confusion matrices. '
    'The broken model (left) predicts Pass for every sample. '
    'The corrected model with optimised threshold (right) detects 6 out of 31 Fail cases.',
    width_in=5.8)

add_body(doc,
    'The correction targets the cause: the model is now penalised proportionally more for '
    'missing a Fail. ROC-AUC is unchanged because it is a ranking metric unaffected by the '
    'threshold. The recall gain from 0% to 19.4% is meaningful in a manufacturing context '
    'where undetected defects have high cost, but it also shows that class weighting alone '
    'is insufficient — precision drops to 20%, producing 4 false alarms for each real defect.',
)

# ── SECTION 7 — BONUS B ───────────────────────────────────────────────────────
add_heading(doc, '7. Failure Mode B: Distribution Shift Between Industrial Phases', level=1, border_bottom=True)

add_heading(doc, 'Symptom and Hypothesis', level=2)
add_body(doc,
    'The SECOM dataset contains six manufacturing phases (Gravure, Lithographie, Métallisation, '
    'Dépôt de couches minces, Implantation ionique, Test électrique). '
    'If a model is trained exclusively on data from five phases and deployed on the sixth, '
    'the feature distributions may differ enough to degrade performance.',
)
add_body(doc,
    'The holdout phase is Implantation ionique, selected because it contains the most '
    'Fail examples (29 out of 261 samples, 11.1%), which makes the evaluation more reliable.',
)

add_heading(doc, 'Controlled Experiment', level=2)
add_body(doc,
    'Broken model: trained on the five source phases; evaluated on both source and target. '
    'Corrected model: includes a calibration subset of the target phase in training. '
    'The evaluation follows a covariate-shift protocol: source performance should remain '
    'stable, and target performance should improve.',
)

add_table(doc,
    headers=['Model', 'Evaluated on', 'Balanced Acc.', 'Recall (Fail)', 'F1 (Fail)', 'ROC-AUC'],
    rows=[
        ['Broken',   'Source phases (test)',   '0.5000', '0.0000', '0.0000', '0.7463'],
        ['Broken',   'Target phase (test)',    '0.5000', '0.0000', '0.0000', '0.7592'],
        ['Corrected','Source phases (test)',   '0.5000', '0.0000', '0.0000', '0.8110'],
        ['Corrected','Target phase (test)',    '0.5000', '0.0000', '0.0000', '0.7328'],
    ],
    col_widths_in=[1.2, 1.85, 1.05, 1.1, 0.9, 0.85],
    row_colors=[BROKEN_BG, BROKEN_BG, FIXED_BG, FIXED_BG],
    bold_col=5,
)
add_note(doc, 'Table 7. Distribution shift results by phase. The effect on the target phase is modest — '
              'the shift between phases in SECOM is weaker than expected.')

add_figure(doc, 'fig7_distribution_shift.png',
    'Figure 7. ROC-AUC by model and evaluation domain. '
    'Adding target-phase data to training improves source ROC-AUC (+0.065) '
    'but slightly reduces target ROC-AUC (−0.026). '
    'The shift effect in SECOM is not as clear-cut as the shortcut experiment.',
    width_in=5.5)

add_body(doc,
    'Discussion: the distribution shift effect is weaker than the shortcut collapse. '
    'Two factors explain this. First, the Phase one-hot encoding is already included as a '
    'feature in the model, so the model is partially aware of phase identity. '
    'Second, the genuine SECOM signal (sensor measurements) may vary more within phases '
    'than between phases. This makes SECOM a clearer demonstration of shortcut learning '
    'than of distribution shift.',
)

# ── SECTION 8 ─────────────────────────────────────────────────────────────────
add_page_break(doc)
add_heading(doc, '8. Threats to Validity', level=1, border_bottom=True)

add_body(doc,
    'Scientific practice requires discussing why the conclusions might be wrong. '
    'We address the five questions required by the project specification:',
)

add_heading(doc, '1. Could the gain be explained by random variability across seeds?', level=2)
add_body(doc,
    'Unlikely. The multi-seed validation on five independent splits shows a consistent gap '
    'of more than 0.95 ROC-AUC points between broken_id and broken_ood, with a standard '
    'deviation of 0.016 on broken_ood. The probability that this gap is due to chance '
    'across all five seeds is negligible. The corrected model also consistently returns '
    'to the reference level (0.77–0.80) across all seeds.',
)

add_heading(doc, '2. Is the test set large enough for the differences to be meaningful?', level=2)
add_body(doc,
    'Partially. The test set contains 471 observations, but only 31 are Fail. '
    'This makes F1 and recall unstable — a single prediction change can shift them '
    'considerably — which is why we use ROC-AUC as the primary metric for the shortcut '
    'experiment. ROC-AUC is computed on 471 probability scores and is more stable. '
    'However, we do not report confidence intervals for ROC-AUC; for a production '
    'deployment, bootstrapped confidence intervals would be required.',
)

add_heading(doc, '3. Could the correction work for reasons unrelated to the identified cause?', level=2)
add_body(doc,
    'This risk is low. The corrected model differs from the broken model by exactly one '
    'change: the removal of shortcut_alarm. The corrected ROC-AUC (0.796) matches the '
    'reference model (0.796) trained on the original features without any shortcut. '
    'If the correction were improving performance via an unrelated mechanism, we would '
    'expect the corrected model to outperform the reference — it does not.',
)

add_heading(doc, '4. Does the experimental protocol generalise beyond SECOM?', level=2)
add_body(doc,
    'The shortcut mechanism is general by construction: the experiment injects a synthetic '
    'shortcut at a controlled reliability level, so the conclusion holds for any dataset '
    'and any reliability value by mathematical necessity. What is dataset-specific is the '
    'magnitude of the effect: with a stronger genuine signal, the model might rely less on '
    'the shortcut. The SECOM signal is weak (reference ROC-AUC ≈ 0.80), which amplifies '
    'the shortcut effect. In a dataset with ROC-AUC ≈ 0.98 without the shortcut, '
    'the collapse on OOD might be less extreme.',
)

add_heading(doc, '5. Is there data leakage that would invalidate the comparison?', level=2)
add_body(doc,
    'No leakage is introduced by the experimental design. '
    'Median imputation is fitted only on the training set, inside the Pipeline, '
    'so test statistics never influence training. '
    'The shortcut feature is added after the train/test split: the train shortcut is '
    'generated from train labels, and each test shortcut variant is generated '
    'independently from test labels with a different seed. '
    'The threshold optimisation for the class imbalance experiment is done on a held-out '
    'inner validation set, not on the final test set.',
)

add_heading(doc, 'Additional Limitations', level=2)
add_body(doc,
    'The shortcut is artificially injected. This proves the mechanism in a controlled '
    'setting but does not demonstrate that a shortcut of this type exists natively in SECOM. '
    'Discovering a naturally occurring shortcut would require domain knowledge from '
    'semiconductor process engineers.',
)
add_body(doc,
    'The analysis uses a single model family (Random Forest). '
    'A different non-linear model (e.g. gradient boosting, MLP) would likely show the same '
    'shortcut collapse, but the specific ROC-AUC values and the sensitivity curve shape '
    'could differ.',
)

# ── SECTION 9 ─────────────────────────────────────────────────────────────────
add_heading(doc, '9. Conclusion and What We Learned', level=1, border_bottom=True)

add_heading(doc, 'Main Findings', level=2)
add_body(doc,
    'The experiment confirms the research hypothesis: a Random Forest trained on SECOM '
    'learns to rely almost exclusively on shortcut_alarm when the feature is 98% reliable '
    'in training. When the correlation is reversed, ROC-AUC collapses from 0.993 to 0.019 — '
    'a near-complete inversion of ranking ability. Removing the shortcut restores performance '
    'to the reference level. The effect is monotonically stronger at higher reliability '
    'values and is consistent across five independent seeds.',
)
add_body(doc,
    'The bonus failure modes show complementary lessons: class imbalance silences the '
    'minority class completely unless the model is explicitly penalised for that silence, '
    'and distribution shift across industrial phases in SECOM is weaker than the shortcut '
    'effect but detectable via ROC-AUC decomposition.',
)

add_heading(doc, 'What We Learned', level=2)
add_body(doc,
    'This project changed how we think about model evaluation:',
)
add_bullet(doc,
    'Accuracy is dangerous on imbalanced data. A model that predicts only the majority '
    'class achieves 93.4% accuracy and appears strong by standard metrics. '
    'Checking recall and ROC-AUC is non-negotiable for rare-event classification.')
add_bullet(doc,
    'A confusion matrix is a symptom; an experiment is a diagnosis. We learned to ask '
    '"what feature is driving this?" rather than "how can I improve the score?" — '
    'and to engineer controlled tests that isolate the cause.')
add_bullet(doc,
    'The deployment environment must be represented in the evaluation. In-distribution '
    'accuracy is not a proxy for robustness. A model can be simultaneously excellent on '
    'the benchmark and useless once deployed if the benchmark does not reflect the '
    'distribution shift it will encounter.')
add_bullet(doc,
    'Feature importance is a tool, not just a result. Discovering that shortcut_alarm '
    'dominated all 452 sensor features was the key diagnostic step. Without this check, '
    'the high ROC-AUC would have been accepted uncritically.')

add_separator(doc)
add_note(doc,
    'Code repository: notebooks/MiniProject_When_ML_Fails_SECOM_BANGOH_DANHOUEGNON_KOUYELE.ipynb  '
    '|  Fixed seed: RANDOM_SEED = 42  '
    '|  Broken pipeline: run_broken_pipeline(...)  '
    '|  Corrected pipeline: run_corrected_pipeline(...)')

# — GitHub link
gh_para = doc.add_paragraph()
gh_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
gh_para.paragraph_format.space_before = Pt(10)
gh_para.paragraph_format.space_after  = Pt(4)
gh_label = gh_para.add_run('GitHub repository : ')
set_run_font(gh_label, size_pt=10, bold=True, color=DARK_BLUE)
gh_link = gh_para.add_run(
    'https://github.com/maounan-018/MiniProject_When_ML_Fails_SECOM_BANGOH_DANHOUEGNON_KOUYELE')
set_run_font(gh_link, size_pt=10, color=BLUE, italic=True)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f'Report saved to: {OUT}')
print(f'File size: {OUT.stat().st_size / 1024:.1f} KB')
