"""Generate the complete project report in French as a separate .docx file."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE    = Path(r'C:\DAVID DANHOUEGNON ECC\OneDrive\DOC DE COURS 2A\Machine Learning\MiniProject_When_ML_Fails_SECOM')
FIGS    = BASE / 'figures'
REPORTS = BASE / 'reports'
REPORTS.mkdir(exist_ok=True)
OUT     = REPORTS / 'MiniProject_When_ML_Fails_SECOM_BANGOH_DANHOUEGNON_KOUYELE_FR.docx'

# ── Colors ────────────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x2E, 0x75, 0xB6)
DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)
ORANGE     = RGBColor(0xC5, 0x5A, 0x11)
GREEN_DARK = RGBColor(0x37, 0x56, 0x23)
TH_BG      = 'D5E8F0'
BROKEN_BG  = 'FCE4D6'
FIXED_BG   = 'E2EFDA'
REF_BG     = 'EBF3FB'

# ── Helpers (identiques à la version EN) ─────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='BFBFBF', size='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'bottom', 'left', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_para_border_bottom(para, color='2E75B6', size='12'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_run_font(run, name='Calibri', size_pt=11, bold=False,
                 italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1, color=None, border_bottom=False):
    para = doc.add_paragraph()
    para.style = doc.styles[f'Heading {level}']
    run = para.add_run(text)
    run.font.name  = 'Calibri'
    run.font.bold  = True
    run.font.size  = Pt({1: 16, 2: 13, 3: 12}.get(level, 11))
    run.font.color.rgb = (color if color else (DARK_BLUE if level == 1 else BLUE))
    if border_bottom:
        set_para_border_bottom(para)
    para.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    para.paragraph_format.space_after  = Pt(6)
    return para

def add_body(doc, text, bold_parts=None, italic=False, indent=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(5)
    para.paragraph_format.space_before = Pt(0)
    if indent:
        para.paragraph_format.left_indent = Cm(0.8)
    if bold_parts is None:
        run = para.add_run(text)
        set_run_font(run, italic=italic)
        return para
    remaining = text
    for chunk, is_bold in bold_parts:
        idx = remaining.find(chunk)
        if idx == -1:
            continue
        if idx > 0:
            r = para.add_run(remaining[:idx])
            set_run_font(r)
        r = para.add_run(chunk)
        set_run_font(r, bold=is_bold, italic=italic)
        remaining = remaining[idx + len(chunk):]
    if remaining:
        r = para.add_run(remaining)
        set_run_font(r)
    return para

def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    run  = para.add_run(text)
    set_run_font(run)
    para.paragraph_format.space_after  = Pt(3)
    para.paragraph_format.space_before = Pt(0)
    if level > 0:
        para.paragraph_format.left_indent = Cm(level * 0.8)
    return para

def add_note(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    set_run_font(run, size_pt=9.5, italic=True,
                 color=RGBColor(0x60, 0x60, 0x60))
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.space_before = Pt(2)
    return para

def add_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run  = para.add_run(text)
    set_run_font(run, size_pt=10, italic=True,
                 color=RGBColor(0x40, 0x40, 0x40))
    para.paragraph_format.space_after  = Pt(10)
    para.paragraph_format.space_before = Pt(2)
    return para

def add_figure(doc, fname, caption, width_in=5.8):
    path = FIGS / fname
    if not path.exists():
        add_note(doc, f'[Figure non trouvée : {fname}]')
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.add_run().add_picture(str(path), width=Inches(width_in))
    add_caption(doc, caption)

def add_table(doc, headers, rows, col_widths_in,
              row_colors=None, bold_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hrow = table.rows[0]
    for i, (cell, hdr) in enumerate(zip(hrow.cells, headers)):
        cell.width = Inches(col_widths_in[i])
        set_cell_bg(cell, TH_BG)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        set_run_font(run, bold=True, size_pt=10, color=DARK_BLUE)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = (row_colors[ri] if row_colors and ri < len(row_colors) else None)
        for ci, (cell, val) in enumerate(zip(row.cells, row_data)):
            cell.width = Inches(col_widths_in[ci])
            if bg:
                set_cell_bg(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if ci == 0
                           else WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run(str(val))
            set_run_font(run, bold=(bold_col is not None and ci == bold_col),
                         size_pt=10)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_separator(doc):
    para = doc.add_paragraph()
    set_para_border_bottom(para, color='D5E8F0', size='6')
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(8)

def add_page_break(doc):
    para = doc.add_paragraph()
    para.add_run().add_break(WD_BREAK.PAGE)

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# En-tête
hdr_para = doc.sections[0].header.paragraphs[0]
hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hdr_run = hdr_para.add_run(
    'ECC Spring 2026  |  Mini-Project: When Machine Learning Fails')
set_run_font(hdr_run, size_pt=9, italic=True,
             color=RGBColor(0x80, 0x80, 0x80))

# Pied de page avec numéros
ftr_para = doc.sections[0].footer.paragraphs[0]
ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
ftr_run  = ftr_para.add_run()
set_run_font(ftr_run, size_pt=9, color=RGBColor(0x80, 0x80, 0x80))
for tag, text in [('begin', None), (None, ' PAGE '), ('end', None)]:
    if tag:
        fc = OxmlElement('w:fldChar')
        fc.set(qn('w:fldCharType'), tag)
        ftr_run._r.append(fc)
    else:
        it = OxmlElement('w:instrText')
        it.text = text
        ftr_run._r.append(it)

# ── PAGE DE GARDE ─────────────────────────────────────────────────────────────
def cover_line(text, size=11, bold=False, italic=False,
               color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size_pt=size, bold=bold, italic=italic,
                 color=color or RGBColor(0x40, 0x40, 0x40))
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def cover_label(text):
    """Petite étiquette de section en majuscules sur la page de garde."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text.upper())
    set_run_font(r, size_pt=10, bold=True,
                 color=RGBColor(0x70, 0x70, 0x70))
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after  = Pt(6)
    return p

# — Informations en haut
cover_line('Centrale Casablanca — ECC Spring 2026',
           size=13, italic=True, color=RGBColor(0x55, 0x55, 0x55),
           space_before=10, space_after=4)
cover_line('Introduction to AI and Machine Learning',
           size=13, italic=True, color=RGBColor(0x55, 0x55, 0x55),
           space_before=0, space_after=0)

# — Grand espace vertical avant le titre
for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(22)

# — Titre principal
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run('When Machine Learning Fails')
r_title.font.name = 'Calibri'
r_title.font.size = Pt(46)
r_title.font.bold = True
r_title.font.color.rgb = DARK_BLUE
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(16)

# — Sous-titre
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run(
    'Diagnosing and Repairing Failure Modes of a Non-Linear Model')
set_run_font(r_sub, size_pt=20, italic=True, color=BLUE)
p_sub.paragraph_format.space_after = Pt(12)
set_para_border_bottom(p_sub, color='2E75B6', size='12')

# — Nature du document
cover_line('Rapport de Mini-Projet', size=15,
           color=RGBColor(0x55, 0x55, 0x55), space_before=10, space_after=5)
cover_line('Dataset SECOM  ·  RandomForestClassifier',
           size=14, italic=True,
           color=RGBColor(0x45, 0x45, 0x45), space_before=0, space_after=0)

# — Auteurs
cover_label('Auteurs')
cover_line('David Danhouegnon  ·  Levis Bangoh  ·  Kouyele Samaa Ashley',
           size=20, bold=True, color=DARK_BLUE,
           space_before=0, space_after=0)

# — Encadrantes
cover_label('Encadrantes')
cover_line('Kawtar Zerhouni  &  Rym Nassih',
           size=18, bold=False, color=BLUE,
           space_before=0, space_after=5)
cover_line('UTER MID@S', size=13, italic=True,
           color=RGBColor(0x70, 0x70, 0x70),
           space_before=0, space_after=0)

add_page_break(doc)

# ── SOMMAIRE ──────────────────────────────────────────────────────────────────
toc_h = doc.add_paragraph()
toc_h.style = doc.styles['Heading 1']
toc_hr = toc_h.add_run('Sommaire')
toc_hr.font.name = 'Calibri'; toc_hr.font.bold = True
toc_hr.font.size = Pt(16); toc_hr.font.color.rgb = DARK_BLUE
set_para_border_bottom(toc_h)
toc_h.paragraph_format.space_before = Pt(4)
toc_h.paragraph_format.space_after  = Pt(12)

toc_para = doc.add_paragraph()
toc_run  = toc_para.add_run()
_fc1 = OxmlElement('w:fldChar');  _fc1.set(qn('w:fldCharType'), 'begin')
_it  = OxmlElement('w:instrText'); _it.set(qn('xml:space'), 'preserve')
_it.text = 'TOC \\o "1-2" \\h \\z \\u'
_fc2 = OxmlElement('w:fldChar');  _fc2.set(qn('w:fldCharType'), 'separate')
_fc3 = OxmlElement('w:fldChar');  _fc3.set(qn('w:fldCharType'), 'end')
toc_run._r.append(_fc1); toc_run._r.append(_it)
toc_run._r.append(_fc2); toc_run._r.append(_fc3)

add_note(doc,
    'Faire un clic droit sur ce champ → « Mettre à jour les champs » '
    'pour afficher le sommaire.')
add_page_break(doc)

# ── SECTION 1 ─────────────────────────────────────────────────────────────────
add_heading(doc, '1. Research Question and Chosen Dataset', level=1,
            border_bottom=True)

add_body(doc,
    'Le dataset SECOM enregistre 590 mesures de capteurs provenant d\'une usine de '
    'fabrication de semi-conducteurs. Chaque observation correspond à un cycle de '
    'production, étiqueté Pass (−1) ou Fail (1). Ce dataset est particulièrement '
    'adapté à ce projet car il est très haute dimension, fortement déséquilibré et '
    'contient de nombreuses valeurs manquantes — des conditions qui permettent '
    'facilement à un modèle de cacher ses défaillances derrière une accuracy globale '
    'apparemment élevée.',
)

add_body(doc, 'Question de recherche :',
         bold_parts=[('Question de recherche :', True)])

rq = doc.add_paragraph()
rq.paragraph_format.left_indent  = Cm(1.0)
rq.paragraph_format.right_indent = Cm(1.0)
rq.paragraph_format.space_before = Pt(4)
rq.paragraph_format.space_after  = Pt(10)
rq_run = rq.add_run(
    '« Un Random Forest entraîné sur SECOM peut-il s\'appuyer sur une variable '
    'parasite non causale, fortement corrélée à la cible pendant l\'entraînement, '
    'au point que sa performance de classement (ROC-AUC) s\'effondre lorsque cette '
    'corrélation est inversée au moment du déploiement ? »'
)
set_run_font(rq_run, italic=True, size_pt=11, color=DARK_BLUE)

add_body(doc,
    'Cette question est falsifiable. L\'hypothèse est rejetée si : '
    '(a) l\'importance de shortcut_alarm reste comparable à celle des vraies variables '
    'capteurs, ou (b) le ROC-AUC ne se dégrade pas significativement lorsque le '
    'shortcut est inversé.',
)

add_heading(doc, 'Dataset and Preprocessing', level=2)
add_body(doc,
    'L\'horodatage est supprimé. Les variables avec plus de 50 % de valeurs manquantes '
    'sont retirées. La colonne catégorielle Phase est encodée en one-hot. Les colonnes '
    'constantes sont supprimées. L\'imputation par la médiane est intégrée dans le '
    'Pipeline scikit-learn afin d\'éviter tout data leakage. '
    'Le découpage train/test (70/30) est stratifié sur la variable cible.',
)

add_table(doc,
    headers=['Propriété', 'Valeur'],
    rows=[
        ['Nombre d\'observations',           '1 567'],
        ['Variables brutes',                 '593'],
        ['Nombre de Fail / taux',            '104 / 6,64 %'],
        ['Variables retirées (>50% manq.)',  '28'],
        ['Variables retirées (constantes)',  '116'],
        ['Variables après nettoyage',        '452'],
        ['Lignes d\'entraînement',           '1 096  (73 Fail — 6,66 %)'],
        ['Lignes de test',                   '471    (31 Fail — 6,58 %)'],
    ],
    col_widths_in=[3.2, 3.0],
)
add_note(doc, 'Tableau 1. Résumé du dataset et du pré-traitement.')

# ── SECTION 2 ─────────────────────────────────────────────────────────────────
add_heading(doc, '2. Reference Model and Observed Symptom', level=1,
            border_bottom=True)

add_body(doc,
    'Le modèle de référence est un Pipeline scikit-learn composé d\'une imputation par '
    'la médiane suivie d\'un RandomForestClassifier avec class_weight="balanced_subsample", '
    'min_samples_leaf=2, 400 arbres et random_state=42. '
    'Le découpage train/test stratifié utilise la même graine fixe.',
)

add_heading(doc, 'Observed Symptom', level=2)
add_body(doc,
    'Sur l\'ensemble de test, l\'accuracy atteint 93,4 % — mais le modèle prédit Pass '
    'pour chaque observation sans exception. Le recall et le F1 sur la classe Fail '
    'sont tous deux nuls. Il s\'agit du piège classique du déséquilibre de classes : '
    'le classificateur de la classe majoritaire est déjà correct 93,4 % du temps, '
    'donc l\'accuracy est une métrique trompeuse pour ce problème.',
)
add_body(doc,
    'Le ROC-AUC = 0,796 sur l\'ensemble de test confirme que le modèle possède une '
    'certaine capacité de classement — il sépare Fail de Pass mieux que le hasard — '
    'mais le seuil de décision par défaut (0,5) ne se déclenche jamais sur la classe '
    'minoritaire. Cette asymétrie justifie le choix du ROC-AUC comme métrique principale '
    'pour l\'expérience sur le shortcut.',
)

add_table(doc,
    headers=['Contexte', 'Accuracy', 'Balanced Acc.', 'Recall (Fail)', 'F1 (Fail)', 'ROC-AUC'],
    rows=[
        ['Référence — train clean', '1,0000', '1,0000', '1,0000', '1,0000', '1,0000'],
        ['Référence — test clean',  '0,9342', '0,5000', '0,0000', '0,0000', '0,7958'],
    ],
    col_widths_in=[2.4, 0.95, 1.1, 1.15, 0.9, 0.9],
    row_colors=[REF_BG, REF_BG],
)
add_note(doc,
    'Tableau 2. Résultats du pipeline de référence. '
    'Le modèle ne prédit aucun Fail sur l\'ensemble de test.')

# ── SECTION 3 ─────────────────────────────────────────────────────────────────
add_heading(doc, '3. Causal Hypothesis and Controlled Experiment', level=1,
            border_bottom=True)

add_heading(doc, 'Causal Hypothesis', level=2)
add_body(doc,
    'Hypothèse : un modèle non linéaire suffisamment flexible s\'appuiera '
    'préférentiellement sur toute variable qui est fortement prédictive de la cible '
    'dans les données d\'entraînement, indépendamment du fait que cette variable soit '
    'causalement liée à l\'issue. Lorsque sa valeur prédictive est supprimée au moment '
    'du test, la performance apparente du modèle s\'effondrera proportionnellement au '
    'degré de dépendance envers cette variable.',
)
add_body(doc,
    'Critère de falsification : si le modèle n\'attribue pas une importance '
    'disproportionnée à la variable parasite, ou si la performance reste stable lorsque '
    'le shortcut est inversé, l\'hypothèse est rejetée.',
)

add_heading(doc, 'Controlled Variable: shortcut_alarm', level=2)
add_body(doc,
    'Une variable binaire synthétique shortcut_alarm est créée à partir des vraies '
    'étiquettes. À l\'entraînement, elle reproduit la cible avec une fiabilité de 98 % '
    '(2 % de retournements aléatoires). Au moment de l\'évaluation, les mêmes '
    'observations de test sont présentées dans trois conditions — '
    'seule la relation shortcut–cible change, tout le reste est identique :',
)
add_bullet(doc,
    'Aligné (ID) :       shortcut_alarm ≈ cible   (corrélation ≈ +0,88)')
add_bullet(doc,
    'Aléatoire :         shortcut_alarm est aléatoire  (corrélation ≈  0,00)')
add_bullet(doc,
    'Inversé (OOD) :    shortcut_alarm ≈ ¬cible  (corrélation ≈ −0,87)')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_table(doc,
    headers=['Condition', 'Corrélation avec la cible'],
    rows=[
        ['Train — aligné',    '+0,8825'],
        ['Test  — aligné',    '+0,8834'],
        ['Test  — aléatoire', '+0,0046'],
        ['Test  — inversé',   '−0,8713'],
    ],
    col_widths_in=[2.8, 2.2],
    row_colors=[FIXED_BG, None, None, BROKEN_BG],
)
add_note(doc,
    'Tableau 3. Diagnostic du shortcut contrôlé. '
    'Les lignes, les étiquettes, la classe de modèle et la graine sont identiques '
    'dans toutes les conditions.')

add_heading(doc, 'Result: ROC-AUC Collapses when Shortcut is Inverted', level=2)
add_body(doc,
    'Le modèle cassé — entraîné avec le shortcut aligné — atteint un ROC-AUC de 0,993 '
    'lorsqu\'il est évalué sur l\'ensemble de test aligné. Lorsque le même modèle est '
    'évalué sur l\'ensemble de test inversé (mêmes lignes, mêmes étiquettes, seul le '
    'shortcut est retourné), le ROC-AUC chute à 0,019. '
    'Ce résultat est pire que l\'aléatoire, ce qui confirme que le modèle a appris à '
    'dépendre presque exclusivement de shortcut_alarm plutôt que des 452 vraies '
    'variables capteurs.',
)

add_figure(doc, 'fig4_confusion_shortcut.png',
    'Figure 1. Confusion matrices on the same test rows: shortcut aligned (left) vs inverted (right). '
    'The confusion pattern reverses almost completely when the shortcut is inverted.',
    width_in=5.8)

add_heading(doc, 'Sensitivity Analysis: Effect Grows with Shortcut Reliability', level=2)
add_body(doc,
    'Pour confirmer que l\'effet est proportionnel à la valeur prédictive apparente du '
    'shortcut, nous faisons varier le paramètre de fiabilité sur quatre niveaux '
    '(0,55 — 0,70 — 0,85 — 0,98) et entraînons un modèle distinct à chaque niveau. '
    'L\'écart entre la performance ID et OOD s\'élargit de façon monotone avec la '
    'fiabilité, ce qui confirme que la dépendance du modèle envers le shortcut est '
    'directement pilotée par la fiabilité de cette variable pendant l\'entraînement.',
)
add_figure(doc, 'fig2_sensitivity.png',
    'Figure 2. ROC-AUC (left) and F1 (right) as a function of shortcut reliability. '
    'The ID/OOD gap widens monotonically — the more reliable the shortcut in training, '
    'the larger the collapse at deployment.',
    width_in=6.0)

add_heading(doc, 'Feature Importance: shortcut_alarm Dominates', level=2)
add_body(doc,
    'L\'importance native (MDI) et l\'importance par permutation (mesurée par la chute '
    'de ROC-AUC) placent toutes deux shortcut_alarm en tête, et ce avec une marge très '
    'large par rapport à l\'ensemble des 452 variables capteurs. '
    'Cela confirme que la performance apparente du modèle est quasi-entièrement portée '
    'par la variable parasite non causale.',
)
add_figure(doc, 'fig3_feature_importance.png',
    'Figure 3. Feature importance for the broken model (shortcut in orange). '
    'Both native MDI (left) and permutation importance (right) place shortcut_alarm '
    'far above all genuine sensor features.',
    width_in=6.0)

# ── SECTION 4 ─────────────────────────────────────────────────────────────────
add_heading(doc, '4. Proposed Correction and Evaluation', level=1,
            border_bottom=True)

add_heading(doc, 'Correction: Remove the Non-Causal Feature', level=2)
add_body(doc,
    'La correction consiste à supprimer shortcut_alarm avant l\'entraînement du pipeline '
    'corrigé. Cette intervention cible directement la cause identifiée — la dépendance '
    'envers un proxy non causal — et non le symptôme. Il ne suffit pas d\'ajuster le '
    'seuil de décision, de modifier le class_weight ou d\'améliorer le découpage '
    'd\'évaluation, car aucune de ces interventions ne supprime l\'incitation du modèle '
    'à exploiter le shortcut.',
)
add_body(doc,
    'Dans le notebook, les pipelines cassé et corrigé sont exposés sous forme de '
    'fonctions indépendantes run_broken_pipeline(...) et run_corrected_pipeline(...), '
    'toutes deux reproductibles à partir des mêmes données préparées et de la même '
    'graine fixe.',
)

add_heading(doc, 'Before / After Comparison', level=2)
add_body(doc,
    'Après suppression du shortcut, le ROC-AUC remonte de 0,019 (OOD, pipeline cassé) '
    'à 0,796, retrouvant exactement le niveau du modèle de référence propre. '
    'Cela confirme le diagnostic causal : le shortcut était l\'unique source de '
    'l\'avantage de performance apparent. Le gain en robustesse OOD est de +0,777 points '
    'de ROC-AUC.',
)
add_body(doc,
    'Note : le recall et le F1 restent nuls dans toutes les configurations car le '
    'déséquilibre de classes (6,6 % de Fail) fait que le seuil de décision 0,5 ne se '
    'déclenche jamais sur la classe minoritaire. C\'est précisément pour cela que le '
    'ROC-AUC — métrique de classement indépendante du seuil — est la mesure appropriée '
    'pour l\'expérience sur le shortcut.',
    bold_parts=[('Note :', True)],
)

add_table(doc,
    headers=['Contexte', 'Balanced Acc.', 'Recall (Fail)', 'F1 (Fail)', 'ROC-AUC'],
    rows=[
        ['Cassé — ID (shortcut aligné)',   '0,5000', '0,0000', '0,0000', '0,9933'],
        ['Cassé — OOD (shortcut inversé)', '0,5000', '0,0000', '0,0000', '0,0193'],
        ['Corrigé — shortcut retiré',      '0,5000', '0,0000', '0,0000', '0,7958'],
        ['Référence — test propre',        '0,5000', '0,0000', '0,0000', '0,7958'],
    ],
    col_widths_in=[2.7, 1.05, 1.15, 0.9, 0.95],
    row_colors=[BROKEN_BG, BROKEN_BG, FIXED_BG, REF_BG],
    bold_col=4,
)
add_note(doc,
    'Tableau 4. Comparaison avant/après correction. '
    'Le ROC-AUC est la métrique clé (sans seuil). '
    'Le modèle corrigé et la référence sont identiques car shortcut_alarm '
    'était la seule variable ajoutée.')

add_figure(doc, 'fig1_roc_comparison.png',
    'Figure 4. ROC-AUC across evaluation conditions. '
    'The broken model collapses from 0.993 to 0.019 when the shortcut is inverted. '
    'The correction restores performance to the reference level (0.796).',
    width_in=5.5)

# ── SECTION 5 ─────────────────────────────────────────────────────────────────
add_heading(doc, '5. Experimental Validation — Variance Across Seeds', level=1,
            border_bottom=True)

add_body(doc,
    'Un seul découpage train/test pourrait donner une conclusion fragile, notamment avec '
    'seulement 31 observations Fail dans l\'ensemble de test. Nous répétons donc '
    'l\'expérience complète sur cinq graines aléatoires indépendantes (7, 11, 19, 23, 42) '
    'et rapportons la moyenne ± l\'écart-type.',
)

add_table(doc,
    headers=['Contexte', 'ROC-AUC moyen', 'ROC-AUC éc.-type', 'F1 moyen', 'F1 éc.-type'],
    rows=[
        ['Cassé — ID (shortcut aligné)',   '0,9928', '0,0011', '0,0000', '0,0000'],
        ['Cassé — OOD (shortcut inversé)', '0,0339', '0,0161', '0,0000', '0,0000'],
        ['Corrigé — test propre',          '0,7682', '0,0248', '0,0000', '0,0000'],
    ],
    col_widths_in=[2.7, 1.1, 1.15, 1.0, 1.0],
    row_colors=[BROKEN_BG, BROKEN_BG, FIXED_BG],
    bold_col=1,
)
add_note(doc,
    'Tableau 5. Validation multi-seeds sur 5 découpages indépendants. '
    'L\'écart entre cassé ID et cassé OOD dépasse 0,95 point de ROC-AUC '
    'pour chaque graine.')

add_figure(doc, 'fig5_multiseed.png',
    'Figure 5. ROC-AUC mean ± std across 5 seeds. '
    'The shortcut collapse and the correction gain are both consistent — '
    'the effect is not due to split variability.',
    width_in=5.2)

add_body(doc,
    'Le très faible écart-type du modèle cassé ID (0,0011) confirme que le mécanisme '
    'du shortcut est quasi-déterministe étant donné sa haute fiabilité (0,98). '
    'Le modèle corrigé présente une variance légèrement plus élevée (0,025), '
    'ce qui est attendu : le signal réel de SECOM est faible et bruité, et de '
    'petites variations dans la composition de l\'ensemble d\'entraînement affectent '
    'la capacité de classement.',
)

# ── SECTION 6 — BONUS A ───────────────────────────────────────────────────────
add_page_break(doc)

bonus_para = doc.add_paragraph()
bonus_run  = bonus_para.add_run('  BONUS — Failure Mode A  ')
set_run_font(bonus_run, bold=True, size_pt=9.5,
             color=RGBColor(0x7F, 0x60, 0x00))
bonus_para.paragraph_format.space_after = Pt(2)

add_heading(doc, '6. Failure Mode A: Class Imbalance Failure', level=1,
            border_bottom=True)

add_heading(doc, 'Observed Symptom', level=2)
add_body(doc,
    'Un Random Forest entraîné sans pondération des classes obtient 93,4 % d\'accuracy '
    'sur l\'ensemble de test SECOM tout en prédisant Pass pour chaque observation — '
    'le recall sur la classe Fail est exactement nul. '
    'Une accuracy élevée masque ici un modèle totalement inutile pour la détection '
    'de défauts, ce qui illustre le danger d\'utiliser l\'accuracy comme seule métrique '
    'sur des données déséquilibrées.',
)

add_heading(doc, 'Causal Hypothesis and Controlled Test', level=2)
add_body(doc,
    'Hypothèse : comme la classe Fail ne représente que 6,6 % des observations, '
    'un modèle qui minimise l\'entropie croisée non pondérée apprend à ne jamais prédire '
    'Fail. Le prédicteur constant « tout Pass » atteint déjà 93,4 % d\'accuracy, '
    'donc le modèle n\'a aucune incitation à apprendre la classe minoritaire.',
)
add_body(doc,
    'Test contrôlé : nous entraînons deux modèles sur les mêmes données — l\'un sans '
    'pondération des classes et l\'autre avec class_weight="balanced_subsample". '
    'Nous optimisons en outre le seuil de décision sur un ensemble de validation '
    'interne en utilisant le score F1 comme critère.',
)

add_heading(doc, 'Correction: Class Weighting + Threshold Optimisation', level=2)
add_body(doc,
    'Le pipeline corrigé utilise class_weight="balanced_subsample" (chaque arbre '
    'rééchantillonne l\'ensemble d\'entraînement pour équilibrer les classes) et optimise '
    'le seuil de décision sur un découpage de validation interne (25 %). '
    'Le seuil optimal trouvé sur la validation est 0,134 — bien en dessous du '
    'seuil par défaut 0,5 — ce qui confirme que la calibration du modèle est mauvaise '
    'en présence d\'un fort déséquilibre de classes.',
)

add_table(doc,
    headers=['Contexte', 'Seuil', 'Balanced Acc.', 'Recall (Fail)', 'F1 (Fail)', 'ROC-AUC'],
    rows=[
        ['Cassé — sans class_weight, test',            '0,500', '0,5000', '0,0000', '0,0000', '0,7792'],
        ['Corrigé — balanced weight, seuil 0,5',       '0,500', '0,5000', '0,0000', '0,0000', '0,7921'],
        ['Corrigé — balanced weight, seuil optimisé',  '0,134', '0,5695', '0,1935', '0,1967', '0,7921'],
    ],
    col_widths_in=[2.7, 0.8, 1.05, 1.1, 0.9, 0.85],
    row_colors=[BROKEN_BG, FIXED_BG, FIXED_BG],
    bold_col=3,
)
add_note(doc,
    'Tableau 6. Résultats sur le déséquilibre de classes. '
    'Le ROC-AUC est inchangé par le seuil (métrique sans seuil). '
    'Le recall passe de 0 % à 19,4 % grâce au seuil optimisé.')

add_figure(doc, 'fig6_confusion_imbalance.png',
    'Figure 6. Class imbalance — confusion matrices. '
    'The broken model (left) predicts Pass for every sample. '
    'The corrected model with optimised threshold (right) detects 6 out of 31 Fail cases.',
    width_in=5.8)

add_body(doc,
    'La correction cible bien la cause : le modèle est désormais pénalisé '
    'proportionnellement plus pour chaque Fail manqué. Le ROC-AUC est inchangé car '
    'c\'est une métrique de classement non affectée par le seuil. '
    'Le gain en recall (de 0 % à 19,4 %) est significatif dans un contexte industriel '
    'où un défaut non détecté a un coût élevé, mais il souligne aussi que la pondération '
    'des classes seule ne suffit pas : la précision chute à 20 %, ce qui génère '
    '4 fausses alarmes pour chaque vrai défaut détecté.',
)

# ── SECTION 7 — BONUS B ───────────────────────────────────────────────────────
add_heading(doc, '7. Failure Mode B: Distribution Shift Between Industrial Phases',
            level=1, border_bottom=True)

add_heading(doc, 'Symptom and Hypothesis', level=2)
add_body(doc,
    'Le dataset SECOM contient six phases de fabrication (Gravure, Lithographie, '
    'Métallisation, Dépôt de couches minces, Implantation ionique, Test électrique). '
    'Si un modèle est entraîné exclusivement sur cinq phases et déployé sur la sixième, '
    'les distributions des variables peuvent être suffisamment différentes pour dégrader '
    'la performance.',
)
add_body(doc,
    'La phase retenue hors entraînement est Implantation ionique, sélectionnée car '
    'elle contient le plus grand nombre d\'exemples Fail (29 sur 261 observations, '
    'soit 11,1 %), ce qui rend l\'évaluation plus fiable.',
)

add_heading(doc, 'Controlled Experiment', level=2)
add_body(doc,
    'Modèle cassé : entraîné sur les cinq phases sources ; évalué à la fois sur la '
    'source et sur la cible. '
    'Modèle corrigé : inclut un sous-ensemble de calibration de la phase cible dans '
    'l\'entraînement. L\'évaluation suit un protocole de covariate shift : la performance '
    'sur la source doit rester stable, et la performance sur la cible doit s\'améliorer.',
)

add_table(doc,
    headers=['Modèle', 'Évalué sur', 'Balanced Acc.', 'Recall (Fail)', 'F1 (Fail)', 'ROC-AUC'],
    rows=[
        ['Cassé',   'Phases source (test)',  '0,5000', '0,0000', '0,0000', '0,7463'],
        ['Cassé',   'Phase cible (test)',    '0,5000', '0,0000', '0,0000', '0,7592'],
        ['Corrigé', 'Phases source (test)',  '0,5000', '0,0000', '0,0000', '0,8110'],
        ['Corrigé', 'Phase cible (test)',    '0,5000', '0,0000', '0,0000', '0,7328'],
    ],
    col_widths_in=[1.2, 1.85, 1.05, 1.1, 0.9, 0.85],
    row_colors=[BROKEN_BG, BROKEN_BG, FIXED_BG, FIXED_BG],
    bold_col=5,
)
add_note(doc,
    'Tableau 7. Résultats du distribution shift par phase. '
    'L\'effet sur la phase cible est modeste — '
    'le shift entre phases dans SECOM est moins prononcé qu\'attendu.')

add_figure(doc, 'fig7_distribution_shift.png',
    'Figure 7. ROC-AUC by model and evaluation domain. '
    'Adding target-phase data to training improves source ROC-AUC (+0.065) '
    'but slightly reduces target ROC-AUC (−0.026).',
    width_in=5.5)

add_body(doc,
    'Discussion : l\'effet du distribution shift est plus faible que celui du shortcut. '
    'Deux facteurs l\'expliquent. Premièrement, l\'encodage one-hot de Phase est déjà '
    'inclus comme variable dans le modèle, qui connaît donc partiellement l\'identité '
    'de la phase. Deuxièmement, le signal réel de SECOM (mesures de capteurs) varie '
    'peut-être davantage à l\'intérieur des phases qu\'entre elles. '
    'Cela fait de SECOM une meilleure démonstration du shortcut learning que du '
    'distribution shift.',
)

# ── SECTION 8 ─────────────────────────────────────────────────────────────────
add_page_break(doc)
add_heading(doc, '8. Threats to Validity', level=1, border_bottom=True)

add_body(doc,
    'La démarche scientifique exige de discuter les raisons pour lesquelles nos '
    'conclusions pourraient être erronées. Nous répondons aux cinq questions '
    'imposées par le sujet :',
)

add_heading(doc,
    '1. Le gain observé pourrait-il s\'expliquer par la variabilité aléatoire entre graines ?',
    level=2)
add_body(doc,
    'C\'est peu probable. La validation multi-seeds sur cinq découpages indépendants '
    'montre un écart constant de plus de 0,95 point de ROC-AUC entre broken_id et '
    'broken_ood, avec un écart-type de 0,016 sur broken_ood. '
    'La probabilité que cet écart soit dû au hasard sur les cinq graines est négligeable. '
    'Le modèle corrigé retrouve lui aussi systématiquement le niveau de référence '
    '(0,77–0,80) quelle que soit la graine.',
)

add_heading(doc,
    '2. L\'ensemble de test est-il assez grand pour que les différences soient significatives ?',
    level=2)
add_body(doc,
    'Partiellement. L\'ensemble de test contient 471 observations, dont seulement 31 '
    'Fail. Cela rend le F1 et le recall instables — un seul changement de prédiction '
    'peut les faire varier considérablement — ce qui justifie l\'utilisation du ROC-AUC '
    'comme métrique principale pour l\'expérience sur le shortcut. '
    'Le ROC-AUC est calculé sur 471 scores de probabilité et est plus stable. '
    'Toutefois, nous ne rapportons pas d\'intervalles de confiance pour le ROC-AUC ; '
    'dans un contexte de déploiement réel, des intervalles bootstrap seraient nécessaires.',
)

add_heading(doc,
    '3. La correction pourrait-elle fonctionner pour des raisons sans rapport avec la cause ?',
    level=2)
add_body(doc,
    'Ce risque est faible. Le modèle corrigé ne diffère du modèle cassé que par un seul '
    'changement : la suppression de shortcut_alarm. Le ROC-AUC du modèle corrigé (0,796) '
    'correspond exactement à celui du modèle de référence (0,796) entraîné sans aucun '
    'shortcut. Si la correction améliorait la performance via un mécanisme non lié à la '
    'cause, on s\'attendrait à ce que le modèle corrigé surpasse la référence — '
    'ce n\'est pas le cas.',
)

add_heading(doc,
    '4. Le protocole expérimental se généralise-t-il au-delà de SECOM ?',
    level=2)
add_body(doc,
    'Le mécanisme du shortcut est général par construction : l\'expérience injecte '
    'un shortcut synthétique à un niveau de fiabilité contrôlé, donc la conclusion '
    'est valide pour tout dataset et toute valeur de fiabilité par nécessité mathématique. '
    'Ce qui est propre au dataset, c\'est l\'ampleur de l\'effet : avec un signal réel '
    'plus fort, le modèle pourrait moins dépendre du shortcut. '
    'Le signal SECOM est faible (ROC-AUC de référence ≈ 0,80), ce qui amplifie l\'effet '
    'du shortcut. Sur un dataset avec ROC-AUC ≈ 0,98 sans shortcut, l\'effondrement '
    'OOD serait probablement moins extrême.',
)

add_heading(doc,
    '5. Y a-t-il du data leakage qui invaliderait la comparaison ?',
    level=2)
add_body(doc,
    'Non. L\'imputation par la médiane est ajustée uniquement sur l\'ensemble '
    'd\'entraînement, à l\'intérieur du Pipeline, donc aucune statistique du test '
    'n\'influence l\'entraînement. '
    'La variable shortcut est ajoutée après le découpage train/test : le shortcut '
    'd\'entraînement est généré à partir des étiquettes d\'entraînement, et chaque '
    'variante du shortcut de test est générée indépendamment à partir des étiquettes '
    'de test avec une graine différente. '
    'L\'optimisation du seuil pour l\'expérience sur le déséquilibre de classes est '
    'réalisée sur un ensemble de validation interne, et non sur l\'ensemble de test final.',
)

add_heading(doc, 'Limitations supplémentaires', level=2)
add_body(doc,
    'Le shortcut est injecté artificiellement. Cela prouve le mécanisme dans un cadre '
    'contrôlé, mais ne démontre pas qu\'un shortcut de ce type existe naturellement dans '
    'SECOM. Identifier un shortcut naturel nécessiterait les connaissances métier d\'un '
    'ingénieur en procédés de fabrication de semi-conducteurs.',
)
add_body(doc,
    'L\'analyse utilise une seule famille de modèles (Random Forest). '
    'Un autre modèle non linéaire (gradient boosting, MLP) montrerait vraisemblablement '
    'le même effondrement, mais les valeurs numériques de ROC-AUC et la forme de la '
    'courbe de sensibilité pourraient différer.',
)

# ── SECTION 9 ─────────────────────────────────────────────────────────────────
add_heading(doc, '9. Conclusion and What We Learned', level=1, border_bottom=True)

add_heading(doc, 'Principaux résultats', level=2)
add_body(doc,
    'L\'expérience confirme l\'hypothèse de recherche : un Random Forest entraîné sur '
    'SECOM apprend à dépendre presque exclusivement de shortcut_alarm lorsque cette '
    'variable est fiable à 98 % pendant l\'entraînement. Lorsque la corrélation est '
    'inversée, le ROC-AUC s\'effondre de 0,993 à 0,019 — une quasi-inversion complète '
    'de la capacité de classement. La suppression du shortcut restaure la performance '
    'au niveau de référence. L\'effet est monotoniquement plus fort à des niveaux de '
    'fiabilité plus élevés et est cohérent sur cinq graines indépendantes.',
)
add_body(doc,
    'Les modes d\'échec bonus apportent des leçons complémentaires : le déséquilibre de '
    'classes réduit au silence la classe minoritaire à moins que le modèle ne soit '
    'explicitement pénalisé pour ce silence, et le distribution shift entre phases '
    'industrielles dans SECOM est détectable via la décomposition du ROC-AUC, '
    'bien que l\'effet soit plus faible que celui du shortcut.',
)

add_heading(doc, 'Ce que nous avons appris', level=2)
add_body(doc,
    'Ce projet a changé notre façon d\'appréhender l\'évaluation des modèles :',
)
add_bullet(doc,
    'L\'accuracy est dangereuse sur des données déséquilibrées. Un modèle qui prédit '
    'uniquement la classe majoritaire obtient 93,4 % d\'accuracy et semble performant '
    'selon les métriques standards. Vérifier systématiquement le recall et le ROC-AUC '
    'est indispensable pour la classification d\'événements rares.')
add_bullet(doc,
    'Une matrice de confusion est un symptôme ; une expérience contrôlée est un '
    'diagnostic. Nous avons appris à nous demander « quelle variable pilote ce résultat ? » '
    'plutôt que « comment améliorer le score ? » — et à concevoir des expériences '
    'contrôlées qui isolent la cause.')
add_bullet(doc,
    'L\'environnement de déploiement doit être représenté dans l\'évaluation. '
    'La performance in-distribution n\'est pas un indicateur de robustesse. '
    'Un modèle peut être simultanément excellent sur le benchmark et totalement '
    'inutile une fois déployé si le benchmark ne reflète pas le distribution shift '
    'qu\'il rencontrera en production.')
add_bullet(doc,
    'L\'importance des variables est un outil de diagnostic, pas seulement un résultat. '
    'Découvrir que shortcut_alarm dominait 452 variables capteurs a été l\'étape clé '
    'du diagnostic. Sans cette vérification, le ROC-AUC élevé aurait été accepté '
    'sans remise en question.')

add_separator(doc)
add_note(doc,
    'Dépôt de code : notebooks/MiniProject_When_ML_Fails_SECOM_BANGOH_DANHOUEGNON_KOUYELE.ipynb  '
    '|  Graine fixe : RANDOM_SEED = 42  '
    '|  Pipeline cassé : run_broken_pipeline(...)  '
    '|  Pipeline corrigé : run_corrected_pipeline(...)')

# — Lien GitHub
gh_para = doc.add_paragraph()
gh_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
gh_para.paragraph_format.space_before = Pt(10)
gh_para.paragraph_format.space_after  = Pt(4)
gh_label = gh_para.add_run('Dépôt GitHub : ')
set_run_font(gh_label, size_pt=10, bold=True, color=DARK_BLUE)
gh_link = gh_para.add_run(
    'https://github.com/maounan-018/MiniProject_When_ML_Fails_SECOM_BANGOH_DANHOUEGNON_KOUYELE')
set_run_font(gh_link, size_pt=10, color=BLUE, italic=True)

# ── Sauvegarde ────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f'Rapport FR sauvegardé : {OUT}')
print(f'Taille : {OUT.stat().st_size / 1024:.1f} KB')
